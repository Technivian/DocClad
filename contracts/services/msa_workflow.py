"""MSA Commercial Review Workflow.

Extends the governed drafting loop proven by the DPA workflow path to MSA
without introducing a second architecture. The flow is still:

New Contract cockpit -> governed draft generation -> Workflow/Contract/
DraftDocument/FieldValue/RiskSignal/ApprovalRoute/CommandCenterWorkItem ->
workspace redirect.
"""
from __future__ import annotations

import re
from calendar import monthrange
from datetime import date, timedelta
from io import BytesIO
from typing import Dict, List, Optional

from django.core.files.base import ContentFile
from django.db import transaction
from django.urls import reverse
from django.utils import timezone

from contracts.middleware import log_action
from contracts.models import (
    ApprovalRoute,
    AuditLog,
    ClauseTemplate,
    CommandCenterWorkItem,
    Contract,
    ContractTemplate,
    Deadline,
    DraftDocument,
    Document,
    FieldDefinition,
    FieldValue,
    RiskSignal,
    Workflow,
    WorkflowStep,
    WorkflowTemplate,
)
from contracts.services.contract_templates import MERGE_FIELDS, _format_value
from contracts.services.workflow_execution import materialize_workflow_from_template
from contracts.tenancy import set_organization_on_instance

WORKFLOW_TEMPLATE_NAME = 'MSA Commercial Review Workflow'
TEMPLATE_LABEL = 'Enterprise Services MSA · Netherlands · B2B'
PLAYBOOK_LABEL = 'MSA Commercial Playbook'
from contracts.services.finance_approval_policy import (
    finance_threshold_display,
    finance_threshold_from_field_values,
    get_finance_approval_threshold,
)

FINANCE_APPROVAL_THRESHOLD = int(get_finance_approval_threshold())
STANDARD_PAYMENT_TERM_DAYS = 30

SECTION_ORDER = [
    FieldDefinition.Section.BASIC_DETAILS,
    FieldDefinition.Section.COMMERCIAL_TERMS,
    FieldDefinition.Section.SERVICES_SCOPE,
    FieldDefinition.Section.LEGAL_POSITION,
    FieldDefinition.Section.SMART_QUESTIONS,
]

_TOKEN_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')


def _add_months(value, months: int):
    """Return a calendar-month offset without requiring another dependency."""
    month_index = value.month - 1 + months
    year = value.year + month_index // 12
    month = month_index % 12 + 1
    return value.replace(day=min(value.day, monthrange(year, month)[1]), year=year, month=month)


def _renewal_notice_due_date(cleaned_values: dict):
    """Derive the first renewal notice date from the governed MSA intake."""
    start_date = cleaned_values.get('start_date')
    if isinstance(start_date, str):
        try:
            start_date = date.fromisoformat(start_date)
        except ValueError:
            return None
    initial_term = str(cleaned_values.get('initial_term') or '')
    match = re.search(r'(\d+)\s*months?', initial_term, re.IGNORECASE)
    if start_date is None or match is None:
        return None
    try:
        notice_days = int(cleaned_values.get('termination_notice_period') or 0)
    except (TypeError, ValueError):
        return None
    if notice_days <= 0:
        return None
    return _add_months(start_date, int(match.group(1))) - timedelta(days=notice_days)


def create_msa_renewal_obligation(*, contract: Contract, workflow: Workflow, cleaned_values: dict, user, request=None):
    """Create one auditable, assigned renewal notice obligation when applicable."""
    is_auto_renew = bool(cleaned_values.get('auto_renewal_included')) or (
        str(cleaned_values.get('renewal_type') or '').lower() == 'auto-renew'
    )
    due_date = _renewal_notice_due_date(cleaned_values) if is_auto_renew else None
    if due_date is None:
        return None

    title = f'Renewal notice — {contract.title}'
    deadline, created = Deadline.objects.get_or_create(
        contract=contract,
        title=title,
        defaults={
            'description': (
                f'Give notice by {due_date.isoformat()} or review the auto-renewal terms for '
                f'{contract.title}.'
            ),
            'deadline_type': Deadline.DeadlineType.RENEWAL,
            'auto_generated': True,
            'generation_source': 'INTAKE',
            'priority': Deadline.Priority.HIGH,
            'due_date': due_date,
            'reminder_days': min(max(int(cleaned_values.get('termination_notice_period') or 7), 7), 90),
            'assigned_to': contract.owner or user,
            'created_by': user,
        },
    )
    if created:
        log_action(
            user,
            AuditLog.Action.CREATE,
            'Deadline',
            deadline.pk,
            str(deadline),
            organization=contract.organization,
            request=request,
            event_type='obligation.auto_created',
            changes={
                'event': 'obligation.auto_created',
                'contract_id': contract.pk,
                'workflow_id': workflow.pk,
                'deadline_type': deadline.deadline_type,
                'due_date': due_date.isoformat(),
            },
        )
    return deadline


def create_msa_document_artifact(*, workflow: Workflow, user, artifact_type: str, request=None) -> Document:
    """Persist a genuine DOCX artifact for the generated MSA or its summary."""
    from docx import Document as WordDocument

    if artifact_type not in {'summary', 'word'}:
        raise ValueError('Unsupported MSA document artifact.')
    draft = workflow.draft_documents.filter(is_current=True).order_by('-version').first()
    if draft is None:
        raise ValueError('No generated MSA draft is available to export.')

    word_document = WordDocument()
    if artifact_type == 'summary':
        word_document.add_heading('MSA Review Summary', level=0)
        word_document.add_paragraph(f'Contract: {workflow.contract.title}')
        word_document.add_paragraph(f'Counterparty: {workflow.contract.counterparty or "Not captured"}')
        word_document.add_paragraph(f'Contract value: {workflow.contract.value or "Not captured"} {workflow.contract.currency or ""}'.strip())
        word_document.add_heading('Active risk signals', level=1)
        signals = list(workflow.risk_signals.order_by('-severity', 'detected_at'))
        if signals:
            for signal in signals:
                word_document.add_paragraph(signal.description, style='List Bullet')
        else:
            word_document.add_paragraph('No active risk signals were detected.')
        word_document.add_heading('Approval status', level=1)
        approvals = workflow.contract.approval_requests.select_related('assigned_to').order_by('created_at')
        if approvals:
            for approval in approvals:
                assignee = approval.assigned_to.get_full_name() or approval.assigned_to.username if approval.assigned_to else 'Unassigned'
                word_document.add_paragraph(f'{approval.approval_step}: {approval.get_status_display()} — {assignee}', style='List Bullet')
        else:
            word_document.add_paragraph('No approval requests have been submitted.')
    else:
        word_document.add_heading('Master Services Agreement', level=0)
        for paragraph in (draft.content or '').split('\n\n'):
            if paragraph.strip():
                word_document.add_paragraph(paragraph.strip())

    payload = BytesIO()
    word_document.save(payload)
    suffix = 'summary' if artifact_type == 'summary' else 'draft'
    from contracts.services.document_version_service import create_document_version

    document, _version = create_document_version(
        organization=workflow.organization,
        title=f'{workflow.contract.title} — {suffix.title()}',
        document_type=Document.DocType.MEMO if artifact_type == 'summary' else Document.DocType.CONTRACT,
        status=Document.Status.DRAFT,
        description=f'Generated from MSA workflow {workflow.pk}.',
        contract=workflow.contract,
        uploaded_by=user,
        actor=user,
        source='generated',
        tags=f'msa-workflow:{workflow.pk},{artifact_type}',
        file=ContentFile(
            payload.getvalue(),
            name=_docx_download_filename(contract=workflow.contract, artifact_type=artifact_type),
        ),
        request=request,
        supersede_prior=False,
    )
    log_action(
        user,
        AuditLog.Action.EXPORT,
        'Document',
        document.pk,
        str(document),
        organization=workflow.organization,
        request=request,
        event_type=f'msa.{artifact_type}_exported',
        changes={
            'event': f'msa.{artifact_type}_exported',
            'contract_id': workflow.contract_id,
            'workflow_id': workflow.pk,
            'document_id': document.pk,
        },
    )
    return document


def get_msa_workflow_template() -> Optional[WorkflowTemplate]:
    return (
        WorkflowTemplate.objects
        .filter(contract_type__code='MSA', is_active=True)
        .order_by('-version')
        .first()
    )


def get_msa_contract_template() -> Optional[ContractTemplate]:
    return (
        ContractTemplate.objects
        .filter(contract_type=Contract.ContractType.MSA, is_active=True)
        .order_by('name')
        .first()
    )


def get_field_definitions_by_section(workflow_template: WorkflowTemplate) -> Dict[str, List[FieldDefinition]]:
    if workflow_template is None:
        return {section: [] for section in SECTION_ORDER}
    definitions = list(
        FieldDefinition.objects.filter(workflow_template=workflow_template).order_by('section', 'order')
    )
    grouped = {section: [] for section in SECTION_ORDER}
    for field in definitions:
        grouped.setdefault(field.section, []).append(field)
    return grouped


def get_msa_approval_route(workflow_template: WorkflowTemplate) -> List[ApprovalRoute]:
    if workflow_template is None:
        return []
    return list(ApprovalRoute.objects.filter(workflow_template=workflow_template).order_by('order'))


def get_clause_library_count(organization, contract_type: str) -> int:
    from django.db.models import Q

    if not contract_type:
        return 0
    qs = ClauseTemplate.objects.filter(is_approved=True)
    qs = qs.filter(Q(organization=organization) | Q(organization__isnull=True)) if organization else qs.filter(organization__isnull=True)
    matching = 0
    for clause in qs.only('applicable_contract_types'):
        allowed = [t.strip() for t in (clause.applicable_contract_types or '').split(',') if t.strip()]
        if not allowed or contract_type in allowed:
            matching += 1
    return matching


def _as_moneyish(value):
    return _format_value(value) if value not in (None, '') else ''


def _payment_term_days(value) -> Optional[int]:
    """Extract the first day count from common demo terms such as Net 30."""
    if value in (None, ''):
        return None
    match = re.search(r'(\d+)', str(value))
    if match is None:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def payment_terms_are_nonstandard(value) -> bool:
    days = _payment_term_days(value)
    return days is not None and days != STANDARD_PAYMENT_TERM_DAYS


def _docx_download_filename(*, contract: Contract, artifact_type: str) -> str:
    """Return the Payrollminds MVP naming convention for generated Word files."""
    client = re.sub(r'[^A-Za-z0-9]+', '_', contract.counterparty or 'Client').strip('_') or 'Client'
    generated_on = timezone.localdate().isoformat()
    suffix = '_Summary' if artifact_type == 'summary' else ''
    return f'Payrollminds_{contract.contract_type}_{client}_{generated_on}{suffix}.docx'


def render_msa_live_preview(template_body: Optional[str], field_values_by_key: dict) -> str:
    if not template_body:
        return ''

    values = dict(field_values_by_key or {})
    personal_data = bool(values.get('personal_data_involved') or values.get('services_involve_personal_data'))
    values.setdefault(
        'data_protection_clause',
        'Data Protection Addendum review is required because the services involve processing personal data.'
        if personal_data else
        'No personal data processing terms are currently required under this draft.'
    )
    values.setdefault(
        'liability_cap',
        'fees paid under this Agreement in the preceding twelve (12) months'
    )

    def _replace(match):
        token = match.group(1)
        if token in values:
            return _format_value(values.get(token))
        if token in MERGE_FIELDS:
            alias = MERGE_FIELDS[token]
            if alias in values:
                return _format_value(values.get(alias))
            return match.group(0)
        return match.group(0)

    return _TOKEN_RE.sub(_replace, template_body)


def detect_msa_risk_signals(workflow: Workflow, field_values_by_key: dict) -> List[RiskSignal]:
    signals = []

    finance_trigger, finance_reason, finance_audit = finance_threshold_from_field_values(field_values_by_key)
    if finance_trigger:
        threshold = get_finance_approval_threshold()
        signals.append((
            'finance_approval_required',
            finance_reason or f'Contract value meets or exceeds the finance approval threshold of {threshold:,.0f}.',
            RiskSignal.Severity.HIGH,
        ))

    if payment_terms_are_nonstandard(field_values_by_key.get('payment_terms')):
        signals.append((
            'nonstandard_payment_terms',
            f'Payment terms deviate from the standard Net {STANDARD_PAYMENT_TERM_DAYS} playbook position.',
            RiskSignal.Severity.HIGH,
        ))

    if field_values_by_key.get('liability_cap_nonstandard'):
        signals.append((
            'liability_cap_nonstandard',
            'Liability cap deviates from the standard MSA playbook position.',
            RiskSignal.Severity.HIGH,
        ))

    if field_values_by_key.get('client_paper'):
        signals.append((
            'client_paper_review_required',
            'Client paper is selected; third-party terms require Legal review before approval.',
            RiskSignal.Severity.HIGH,
        ))

    personal_data = bool(field_values_by_key.get('personal_data_involved') or field_values_by_key.get('services_involve_personal_data'))
    if personal_data:
        signals.append((
            'msa_dpa_review_required',
            'Services involve personal data processing; privacy review and linked DPA assessment are required.',
            RiskSignal.Severity.MEDIUM,
        ))

    auto_renew = bool(field_values_by_key.get('auto_renewal_included')) or str(field_values_by_key.get('renewal_type', '')).lower() == 'auto-renew'
    if auto_renew:
        signals.append((
            'renewal_notice_review',
            'Auto-renewal is included; review notice periods and renewal obligation tracking.',
            RiskSignal.Severity.MEDIUM,
        ))

    if field_values_by_key.get('ip_ownership_nonstandard'):
        signals.append((
            'nonstandard_ip_ownership',
            'IP ownership is non-standard and requires Legal review.',
            RiskSignal.Severity.HIGH,
        ))

    governing_law = str(field_values_by_key.get('governing_law') or '').strip().lower()
    governing_law_nonpreferred = bool(field_values_by_key.get('governing_law_nonpreferred')) or (governing_law and 'netherlands' not in governing_law)
    if governing_law_nonpreferred:
        signals.append((
            'nonpreferred_governing_law',
            'Governing law falls outside the preferred jurisdiction and requires Legal escalation.',
            RiskSignal.Severity.MEDIUM,
        ))

    created = []
    for code, description, severity in signals:
        created.append(RiskSignal.objects.create(workflow=workflow, code=code, description=description, severity=severity))
    return created


def sync_command_center_work_item_for_workflow(workflow: Workflow) -> CommandCenterWorkItem:
    contract = workflow.contract
    current_step = (
        workflow.steps.filter(status=WorkflowStep.Status.IN_PROGRESS).order_by('order').first()
        or workflow.steps.filter(status=WorkflowStep.Status.PENDING).order_by('order').first()
    )
    risk_signals = list(workflow.risk_signals.filter(is_resolved=False).order_by('-severity', 'detected_at'))
    severity_rank = {
        RiskSignal.Severity.CRITICAL: 4,
        RiskSignal.Severity.HIGH: 3,
        RiskSignal.Severity.MEDIUM: 2,
        RiskSignal.Severity.LOW: 1,
    }
    top_signal = max(risk_signals, key=lambda signal: severity_rank.get(signal.severity, 0)) if risk_signals else None
    current_stage = current_step.name if current_step else 'Intake'
    open_exceptions = len(risk_signals)

    if open_exceptions:
        next_action = f'Resolve {open_exceptions} exception{"s" if open_exceptions != 1 else ""}'
    elif top_signal and top_signal.code in {'finance_approval_required', 'nonstandard_payment_terms'}:
        next_action = 'Review approval route'
    elif top_signal and top_signal.code == 'msa_dpa_review_required':
        next_action = 'Review privacy scope and linked DPA need'
    elif top_signal and top_signal.code == 'renewal_notice_review':
        next_action = 'Review renewal notice obligations'
    elif top_signal:
        next_action = 'Review legal fallback positions'
    else:
        next_action = 'Open generated MSA workspace'

    blocking_issue = (
        f'{open_exceptions} open exception{"s" if open_exceptions != 1 else ""} must be resolved before review.'
        if open_exceptions
        else (top_signal.description if top_signal else 'Field capture complete; ready for governed review.')
    )
    item, _ = CommandCenterWorkItem.objects.update_or_create(
        organization=workflow.organization,
        source_type=CommandCenterWorkItem.SourceType.WORKFLOW,
        source_model='Workflow',
        source_object_id=workflow.pk,
        defaults={
            'title': workflow.title,
            'subtitle': blocking_issue,
            'item_type': 'MSA workflow',
            'stage': current_stage,
            'status': (
                CommandCenterWorkItem.Status.BLOCKED
                if top_signal and top_signal.severity in {RiskSignal.Severity.HIGH, RiskSignal.Severity.CRITICAL}
                else CommandCenterWorkItem.Status.OPEN
            ),
            'risk_level': top_signal.severity if top_signal else (contract.risk_level if contract else Contract.RiskLevel.LOW),
            'priority': (
                CommandCenterWorkItem.Priority.HIGH
                if top_signal and top_signal.severity in {RiskSignal.Severity.HIGH, RiskSignal.Severity.CRITICAL}
                else CommandCenterWorkItem.Priority.MEDIUM
            ),
            'owner': workflow.created_by,
            'contract': contract,
            'workflow': workflow,
            'due_at': current_step.due_date if current_step else None,
            'action_label': 'Open workspace',
            'action_path': reverse('contracts:workflow_detail', kwargs={'pk': workflow.pk}),
            'flags': {
                'contract_type': 'MSA',
                'current_stage': current_stage,
                'owner_label': (
                    workflow.created_by.get_full_name() or workflow.created_by.username
                    if workflow.created_by else 'Unassigned'
                ),
                'highest_risk_signal': top_signal.description if top_signal else 'No active risk signal',
                'blocking_issue': blocking_issue,
                'next_action': next_action,
                'open_exceptions': open_exceptions,
                'self_serve_eligible': contract.risk_level == Contract.RiskLevel.LOW if contract else False,
            },
            'last_source_synced_at': timezone.now(),
        },
    )
    return item


@transaction.atomic
def create_msa_workflow_instance(*, organization, user, cleaned_values: dict, request=None) -> Workflow:
    workflow_template = get_msa_workflow_template()
    if workflow_template is None:
        raise ValueError('MSA Commercial Review Workflow template is not seeded.')

    field_defs = list(FieldDefinition.objects.filter(workflow_template=workflow_template))

    contract = Contract(
        title=f"MSA — {cleaned_values.get('counterparty') or 'Untitled counterparty'}",
        contract_type=Contract.ContractType.MSA,
        status=Contract.Status.IN_PROGRESS,
        created_by=user,
        lifecycle_stage=Contract.LifecycleStage.DRAFTING,
        risk_level=Contract.RiskLevel.LOW,
    )
    set_organization_on_instance(contract, organization)
    for field in field_defs:
        if field.maps_to_contract_field and field.key in cleaned_values:
            setattr(contract, field.maps_to_contract_field, cleaned_values[field.key])
    contract.auto_renew = bool(cleaned_values.get('auto_renewal_included')) or str(cleaned_values.get('renewal_type', '')).lower() == 'auto-renew'
    from contracts.services.contract_provenance import OriginKind, apply_provenance_fields, pin_workflow_provenance
    apply_provenance_fields(
        contract,
        origin_kind=OriginKind.WORKFLOW,
        origin_channel='msa_workflow',
        actor=user,
        lock=False,
        validate=False,
    )
    contract.save()

    workflow = Workflow.objects.create(
        title=WORKFLOW_TEMPLATE_NAME,
        description=f"Commercial review for the MSA with {cleaned_values.get('counterparty') or 'this counterparty'}.",
        organization=organization,
        template=workflow_template,
        contract=contract,
        status=Workflow.Status.ACTIVE,
        created_by=user,
    )
    pin_workflow_provenance(contract, workflow, actor=user, request=request, channel='msa_workflow')
    materialize_workflow_from_template(workflow)

    FieldValue.objects.bulk_create([
        FieldValue(workflow=workflow, field_definition=field, value=cleaned_values.get(field.key))
        for field in field_defs
    ])

    contract_template = get_msa_contract_template()
    preview_content = render_msa_live_preview(contract_template.body if contract_template else None, cleaned_values)
    DraftDocument.objects.create(
        workflow=workflow,
        contract=contract,
        content=preview_content,
        version=1,
        is_current=True,
        created_by=user,
    )
    log_action(
        user,
        AuditLog.Action.CREATE,
        'Workflow',
        workflow.pk,
        str(workflow),
        organization=organization,
        request=request,
        event_type='msa.template_applied',
        changes={'event': 'msa.template_applied', 'contract_id': contract.pk},
    )
    log_action(
        user,
        AuditLog.Action.CREATE,
        'Workflow',
        workflow.pk,
        str(workflow),
        organization=organization,
        request=request,
        event_type='msa.fields_captured',
        changes={'event': 'msa.fields_captured', 'contract_id': contract.pk, 'field_count': len(field_defs)},
    )

    risk_signals = detect_msa_risk_signals(workflow, cleaned_values)
    if any(signal.severity in {RiskSignal.Severity.HIGH, RiskSignal.Severity.CRITICAL} for signal in risk_signals):
        contract.risk_level = Contract.RiskLevel.HIGH
    elif risk_signals:
        contract.risk_level = Contract.RiskLevel.MEDIUM
    contract.save(update_fields=['risk_level', 'auto_renew', 'updated_at'])
    log_action(
        user,
        AuditLog.Action.CREATE,
        'Workflow',
        workflow.pk,
        str(workflow),
        organization=organization,
        request=request,
        event_type='msa.risks_evaluated',
        changes={
            'event': 'msa.risks_evaluated',
            'contract_id': contract.pk,
            'risk_signal_count': len(risk_signals),
        },
    )
    create_msa_renewal_obligation(
        contract=contract,
        workflow=workflow,
        cleaned_values=cleaned_values,
        user=user,
        request=request,
    )

    sync_command_center_work_item_for_workflow(workflow)

    log_action(
        user,
        'CREATE',
        'Workflow',
        workflow.id,
        str(workflow),
        changes={'event': 'msa_workflow_created', 'contract_id': contract.id},
        request=request,
        organization=organization,
    )
    return workflow
